"""투자심사보고서에서 데이터를 추출하는 모듈.
테이블 위치/구조에 무관하게 키워드 기반으로 자동 탐색한다."""
import re
from dataclasses import dataclass, field
from docx import Document


@dataclass
class InvestmentReportData:
    company_name: str = ""
    representative: str = ""
    address: str = ""
    establishment_date: str = ""
    paid_in_capital: str = ""
    par_value: str = ""
    num_employees: str = ""
    business_description: str = ""
    business_registration: str = ""
    industry_code: str = ""            # 한국표준산업분류코드 (예: C28114)

    review_date: str = ""
    committee_date: str = ""
    contract_date: str = ""
    fund_date: str = ""
    fund_name: str = ""
    discoverer: str = ""
    reviewer: str = ""
    post_manager: str = ""

    investment_amount: str = ""
    issue_price: str = ""
    total_shares: str = ""
    share_ratio: str = ""
    stock_type: str = ""
    pre_value: str = ""
    post_value: str = ""

    duration: str = ""
    redemption_terms: str = ""
    conversion_terms: str = ""
    refixing_terms: str = ""
    other_terms: str = ""
    fund_usage: str = ""

    buyback_rate: str = ""
    penalty_rate: str = ""
    delay_rate: str = ""

    co_investors: list = field(default_factory=list)
    discovery_background: str = ""

    # 별첨2 투자재원검토보고서 (표5 준법사항용)
    investment_type: str = ""          # 투자구분 (신규발행/구주 등)
    is_follow_up: str = ""             # 후속투자 여부
    purpose_transport: str = ""        # 주목적 - 국토교통분야 해당여부
    purpose_mobility: str = ""         # 주목적 - 혁신성장 모빌리티 분야
    purpose_south: str = ""            # 주목적 - 산업은행 남부권 전략산업
    purpose_tcb: str = ""              # 주목적 - TCB Ti-6 등급 이상
    purpose_tcb_detail: str = ""       # TCB 상세 (등급, 발급일)
    purpose_ibk: str = ""              # 주목적 - 중소기업은행 거래기업

    # 벤처기업/이노비즈 인증
    is_venture: str = ""               # 벤처기업 인증 여부
    venture_expiry: str = ""           # 벤처기업확인서 유효기간
    is_innobiz: str = ""              # 이노비즈/메인비즈 인증 여부
    innobiz_expiry: str = ""          # 이노비즈 유효기간

    warnings: list = field(default_factory=list)


# ── 지역 리스트 ──
_REGIONS = (
    '서울특별시|부산광역시|대구광역시|인천광역시|광주광역시|대전광역시|울산광역시|세종특별자치시'
    '|경기도|강원도|강원특별자치도|충청북도|충청남도|전라북도|전북특별자치도|전라남도|경상북도|경상남도|제주특별자치도'
    '|서울|경기|인천|부산|대구|광주|대전|울산|세종|강원|충북|충남|전북|전남|경북|경남|제주'
    '|대전시|부산시'
)


def extract_report_data(filepath: str) -> InvestmentReportData:
    ext = filepath.lower().rsplit('.', 1)[-1] if '.' in filepath else ''
    data = InvestmentReportData()

    if ext == 'pdf':
        from extractors.pdf_extractor import extract_text_from_pdf
        full_text = extract_text_from_pdf(filepath)
        doc = None
        _extract_from_pdf_text(full_text, data)
    else:
        doc = Document(filepath)
        paragraphs = [p.text.strip() for p in doc.paragraphs]
        full_text = "\n".join(paragraphs)
        _scan_all_tables(doc.tables, data)
        # 제목("투자심사보고서") 직후 paragraph에서 회사명 fallback 추출
        _extract_company_name_from_title(paragraphs, data)

    _extract_from_text(full_text, data, doc)

    # 별첨1 주요투자조건 요약서 (우선 참조)
    if doc:
        _extract_appendix1_summary(doc.tables, data)
        # 동반투자 테이블
        if not data.co_investors:
            _extract_co_investors_table(doc.tables, data)

    # 별첨2 투자재원검토보고서 + 벤처/이노비즈 인증
    if doc:
        _extract_appendix2(doc.tables, data)
    _extract_certifications(full_text, data)

    return data


# ━━━━━━━━━━━━━━━ 테이블 자동 탐색 ━━━━━━━━━━━━━━━

def _scan_all_tables(tables, data: InvestmentReportData):
    """모든 테이블을 순회하며 키워드로 테이블 유형을 감지한다."""
    for table in tables:
        if not table.rows:
            continue

        # 테이블 전체 텍스트 (처음 3행)로 유형 판별
        sample_text = ""
        for row in table.rows[:3]:
            sample_text += " ".join(c.text.strip() for c in row.cells) + " "

        # ── 일정/담당자 테이블 ──
        if not data.review_date and ('투심' in sample_text or '일정' in sample_text):
            _extract_schedule_table(table, data)

        # ── 회사 개요 테이블 ──
        if not data.company_name and (
            '회사명' in sample_text or '대표이사' in sample_text
            or '대표자' in sample_text or '납입자본금' in sample_text
        ):
            _extract_company_table(table, data)

        # ── 투자조건 요약 테이블 (별첨1: 투자업체/총투자금액/주요조건/사용용도 등) ──
        # fund_usage 외에도 company_name/투자조건 추출 대상이므로 항상 시도
        if (not data.fund_usage or not data.company_name
                or not data.investment_amount or not data.duration):
            _extract_investment_summary_table(table, data)

    # ── 주주현황 → 지분율 ──
    _extract_shareholder_table(tables, data)


def _extract_schedule_table(table, data: InvestmentReportData):
    """일정/담당자 테이블. 키워드 매칭."""
    _FIELD_MAP = {
        '투심 일자': 'review_date', '투심일자': 'review_date',
        '조합투심위': 'committee_date',
        '계약일': 'contract_date',
        '자금집행일': 'fund_date',
        '투자 계정': 'fund_name', '투자계정': 'fund_name',
        '발굴': 'discoverer',
        '심사': 'reviewer',
        '사후관리': 'post_manager',
    }
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        row_text = " ".join(cells)
        for keyword, attr in _FIELD_MAP.items():
            if keyword in row_text and not getattr(data, attr):
                # 값은 마지막 셀 (또는 키워드 다음 셀)
                val = cells[-1].strip()
                if val and keyword not in val:
                    setattr(data, attr, val)
                else:
                    for i, c in enumerate(cells):
                        if keyword in c and i + 1 < len(cells):
                            setattr(data, attr, cells[i + 1].strip())
                            break


def _extract_company_table(table, data: InvestmentReportData):
    """회사 개요 테이블. 다양한 필드명 지원."""
    _FIELD_MAP = {
        '회사명': 'company_name',
        '대표이사': 'representative', '대표자': 'representative',
        '납입자본금': 'paid_in_capital',
        '액면가': 'par_value',
        '설립일': 'establishment_date',
        '인력현황': 'num_employees', '종업원': 'num_employees', '종업원수': 'num_employees',
        '주요사업': 'business_description', '주요 사업': 'business_description',
    }

    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) < 2:
            continue
        row_text = " ".join(cells)

        # 일반 필드 매칭
        for keyword, attr in _FIELD_MAP.items():
            if keyword in row_text and not getattr(data, attr):
                for i, c in enumerate(cells):
                    if keyword in c and i + 1 < len(cells):
                        val = cells[i + 1].replace('\xa0', '').strip()
                        if attr == 'representative':
                            val = val.replace(' ', '')
                            # 이름만 추출 (학력/경력이 붙어있을 경우)
                            m_name = re.match(r'([가-힣]{2,4})', val)
                            if m_name:
                                val = m_name.group(1)
                        setattr(data, attr, val)
                        break

        # 사업자등록번호 (패턴 매칭)
        if ('사업자' in row_text or '등록번호' in row_text) and not data.business_registration:
            for c in cells:
                m = re.search(r'(\d{3}-\d{2}-\d{5})', c)
                if m:
                    data.business_registration = m.group(1)
                    break

        # 주소 (다양한 표현)
        if not data.address and ('주소' in row_text or '본사' in row_text or '소재지' in row_text):
            for c in reversed(cells):
                if c and '주소' not in c and '본사' not in c and '소재지' not in c and len(c) > 5:
                    data.address = c.replace('\n', ' ').strip()
                    break


def _extract_company_name_from_title(paragraphs: list, data: InvestmentReportData):
    """제목('투자심사보고서') 직후 paragraph에서 회사명 fallback 추출.
    예: paragraph[4]='투자심사보고서', paragraph[5]='㈜듀셀'"""
    if data.company_name:
        return
    for i, p in enumerate(paragraphs):
        if '투자심사보고서' in p or '투자검토보고서' in p:
            # 다음 5개 paragraph에서 회사명 후보 탐색
            for j in range(i + 1, min(i + 6, len(paragraphs))):
                cand = paragraphs[j].strip()
                if not cand or len(cand) > 30:
                    continue
                # 케이런/조합/펀드 등 투자자측 명칭 제외
                if any(kw in cand for kw in ['케이런', '조합', '펀드', '벤처스', '유한회사']):
                    continue
                # 회사명 형식: ㈜XXX, 주식회사 XXX, XXX 주식회사, XXX㈜
                if re.search(r'(?:㈜|주식회사|\(주\))', cand):
                    data.company_name = cand
                    return
            break


def _extract_investment_summary_table(table, data: InvestmentReportData):
    """투자조건 요약 테이블에서 회사명, 투자조건, 사용용도, 위약벌 등 추출.
    별첨1 주요 투자조건 요약서(투자업체/총투자금액/투자방법/투자단가/주요조건 등) 형태 지원."""
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) < 2:
            continue
        label = cells[0]
        value = cells[1] if len(cells) > 1 else ""

        # 회사명 (투자업체)
        if '투자업체' in label and not data.company_name and value:
            v = value.replace('\xa0', '').strip()
            if v and '케이런' not in v and '조합' not in v and '펀드' not in v:
                data.company_name = v

        # 총투자금액
        if ('총투자금액' in label or '투자금액' in label) and not data.investment_amount:
            m = re.search(r'([\d,]{6,})\s*원', value)
            if m:
                data.investment_amount = m.group(1) + "원"

        # 투자단가 (예: "6,214원 (액면가: 500원)")
        if '투자단가' in label and not data.issue_price:
            m = re.search(r'([\d,]+)\s*원', value)
            if m:
                data.issue_price = m.group(1) + "원"
            m2 = re.search(r'액면가\s*[:：]?\s*([\d,]+)\s*원', value)
            if m2 and not data.par_value:
                data.par_value = m2.group(1)

        # 인수주식수
        if '인수주식수' in label and not data.total_shares:
            m = re.search(r'([\d,]+)\s*주', value)
            if m:
                data.total_shares = m.group(1) + "주"

        # 투자전 기업가치 (Pre-Value)
        if '투자전' in label and '기업가치' in label and not data.pre_value:
            m = re.search(r'(\d+)\s*억', value)
            if m:
                data.pre_value = m.group(1) + "억원"

        # 투자방식 (투자방법: 신규발행 RCPS 인수 등)
        if '투자방법' in label and not data.stock_type:
            if 'RCPS' in value:
                data.stock_type = '상환전환우선주'
            elif 'CPS' in value or '전환우선주' in value:
                data.stock_type = '전환우선주'
            elif 'RPS' in value or '상환우선주' in value:
                data.stock_type = '상환우선주'
            elif 'CB' in value or '전환사채' in value:
                data.stock_type = '전환사채'
            elif 'BW' in value or '신주인수권부사채' in value:
                data.stock_type = '신주인수권부사채'
            elif '보통주' in value:
                data.stock_type = '보통주'

        # 주요조건 (존속기간, 상환, Refixing)
        if '주요조건' in label:
            if not data.duration:
                m = re.search(r'존속기간\s*(\d+)\s*년', value)
                if m:
                    data.duration = m.group(1) + "년"
            if not data.redemption_terms:
                m = re.search(r'상환청구.*?(\d+)\s*년\s*후.*?YTM\s*(\d+)\s*%', value)
                if m:
                    data.redemption_terms = f"{m.group(1)}년후부터 상환청구 가능, 연복리 {m.group(2)}%"
                else:
                    m = re.search(r'YTM\s*(\d+)\s*%', value)
                    if m:
                        data.redemption_terms = f"YTM {m.group(1)}%"
            if not data.refixing_terms:
                m = re.search(r'IPO[/]?M&A\s*리?픽?싱?\s*(\d+)\s*%', value)
                if m:
                    data.refixing_terms = f"IPO/M&A {m.group(1)}%"

        # 동반투자내역
        if '동반투자' in label and not data.co_investors:
            for line in value.split('\n'):
                line = line.strip().lstrip('-').strip()
                if line and not line.endswith(')'):
                    # "IBK벤처투자-퓨처플레이" 형태
                    if not any(kw in line for kw in ['Post-Value', 'Pre-Value', '납입', '신주']):
                        data.co_investors.append((line, "", ""))

        # 투자금 사용용도
        if '투자금 사용용도' in label and '위반' not in label and not data.fund_usage:
            data.fund_usage = value

        # 의무불이행 이자율 (보통 단일 % 값)
        if '의무불이행' in label and '이자율' in label:
            m = re.search(r'(\d+)\s*%', value)
            if m:
                rate = m.group(1)
                if not data.penalty_rate:
                    data.penalty_rate = rate
                if not data.delay_rate:
                    data.delay_rate = rate
                if not data.buyback_rate:
                    data.buyback_rate = rate


def _extract_shareholder_table(tables, data: InvestmentReportData):
    """주주현황 테이블에서 케이런 지분율 추출."""
    if data.share_ratio:
        return

    for table in tables:
        total_post_shares = 0
        keiren_row_data = None

        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = " ".join(cells)

            # 합계행
            if '합계' in row_text:
                for c in cells:
                    m = re.match(r'^([\d,]+)$', c.strip())
                    if m:
                        val = int(m.group(1).replace(',', ''))
                        if val > total_post_shares:
                            total_post_shares = val

            # 케이런 행
            if '케이런' in row_text and ('주' in row_text or '%' in row_text or any(c.replace(',','').isdigit() for c in cells)):
                keiren_row_data = cells

        if keiren_row_data:
            # 주식수 추출
            shares_in_row = []
            for c in keiren_row_data:
                m = re.search(r'^([\d,]+)$', c.strip())
                if m:
                    val = int(m.group(1).replace(',', ''))
                    if val > 100:
                        shares_in_row.append(val)

            # 비율 추출
            ratios = []
            for c in keiren_row_data:
                m = re.search(r'(\d+\.?\d*)\s*%', c)
                if m:
                    ratios.append(m.group(1))

            # 직접 계산 우선
            if total_post_shares > 0 and shares_in_row:
                our_shares = max(shares_in_row)
                calculated = round(our_shares / total_post_shares * 100, 2)
                data.share_ratio = f"{calculated}%"
                return
            if ratios:
                data.share_ratio = ratios[-1] + "%"
                return


# ━━━━━━━━━━━━━━━ 본문 텍스트 추출 ━━━━━━━━━━━━━━━

def _extract_from_text(full_text: str, data: InvestmentReportData, doc=None):
    """본문에서 투자조건, 동반투자 등 추출."""

    # ── 투자금액 ──
    if not data.investment_amount:
        for pat in [
            r'총\s*([\d,]{7,})\s*원',
            r'([\d,]{7,})\s*원.*?규모',
            r'투자금액\s*[:：]?\s*([\d,]{7,})\s*원',
        ]:
            m = re.search(pat, full_text)
            if m:
                data.investment_amount = m.group(1) + "원"
                break

    # ── 투자단가 ──
    if not data.issue_price:
        for pat in [
            r'주당\s*([\d,]+)\s*원',
            r'투자단가\s*[:：]?\s*([\d,]+)\s*원',
        ]:
            m = re.search(pat, full_text)
            if m:
                data.issue_price = m.group(1) + "원"
                break

    # ── 주식수 ──
    if not data.total_shares:
        m = re.search(r'([\d,]+)\s*주를?\s*주당', full_text)
        if m:
            data.total_shares = m.group(1) + "주"

    # ── 투자방식 ──
    if not data.stock_type:
        m = re.search(r'(상환전환우선주|전환우선주|상환우선주|보통주)', full_text)
        if m:
            data.stock_type = m.group(1)

    # ── Pre/Post 기업가치 ──
    if not data.pre_value:
        m = re.search(r'Pre[- ]?[Vv]alue.*?(\d+)\s*억', full_text)
        if m:
            data.pre_value = m.group(1) + "억원"
    if not data.post_value:
        m = re.search(r'Post[- ]?[Vv]alue.*?(\d+)\s*억', full_text)
        if m:
            data.post_value = m.group(1) + "억원"

    # ── 존속기간/상환/Refixing ──
    if not data.duration:
        m = re.search(r'만기\s*(\d+)\s*년', full_text)
        if m:
            data.duration = m.group(1) + "년"
    if not data.redemption_terms:
        m = re.search(r'발행\s*(\d+)\s*년\s*후.*?상환.*?(?:YTM|연복리)\s*(\d+)\s*%', full_text)
        if m:
            data.redemption_terms = f"{m.group(1)}년후부터 상환청구 가능, 연복리 {m.group(2)}%"
    if not data.refixing_terms:
        m = re.search(r'IPO/M&A\s*(\d+)\s*%', full_text)
        if m:
            data.refixing_terms = f"IPO/M&A {m.group(1)}%"

    # ── 투자금 사용용도 (테이블 우선, 텍스트 fallback) ──
    if not data.fund_usage and doc:
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 2 and '투자금 사용용도' in cells[0] and '위반' not in cells[0]:
                    data.fund_usage = cells[1].strip()
                    break
            if data.fund_usage:
                break
    if not data.fund_usage:
        for line in full_text.split('\n'):
            if any(kw in line for kw in ['운영자금', '설비투자', '연구개발', '시설자금']):
                if len(line) < 200:
                    data.fund_usage = line.strip()
                    break

    # ── 동반투자 ──
    if not data.co_investors:
        m = re.search(r'동반투자(?:기관|내역)[：:]?\s*(.+)', full_text)
        if m:
            pairs = re.findall(r'(\S+)\s+(\d+)억', m.group(1))
            for name, amount in pairs:
                data.co_investors.append((name, f"{amount}억원", ""))


# ━━━━━━━━━━━━━━━ PDF 전용 ━━━━━━━━━━━━━━━

def _extract_from_pdf_text(full_text: str, data: InvestmentReportData):
    """PDF 텍스트에서 테이블 없이 모든 데이터를 추출."""
    m = re.search(r'[㈜(주)]\s*(\S+)', full_text)
    if m:
        data.company_name = "㈜" + m.group(1)

    m = re.search(r'대표이사[:\s]*([가-힣]{2,4})', full_text)
    if m:
        data.representative = m.group(1)

    m = re.search(r'(\d{3}-\d{2}-\d{5})', full_text)
    if m:
        data.business_registration = m.group(1)

    m = re.search(r'(' + _REGIONS + r')[^\n]{5,80}', full_text)
    if m:
        data.address = m.group(0).strip()

    m = re.search(r'설립일[:\s]*([\d.]+)', full_text)
    if m:
        data.establishment_date = m.group(1)

    for pat in [
        r'((?:케이런|IBK)\S*\s*\S*\s*(?:투자)?조합)',
        r'(\S+\d+호\S*\s*(?:투자)?조합)',
        r'(20\d{2}\s*\S*\s*\S*\s*\d+호\s*펀드)',
    ]:
        m = re.search(pat, full_text)
        if m:
            data.fund_name = m.group(1)
            break

    m = re.search(r'발굴[:\s]*(\S+(?:\s*\(\d+%\))?)', full_text)
    if m:
        data.discoverer = m.group(1)
    m = re.search(r'심사[:\s]*(\S+(?:\s*\(\d+%\))(?:\s*,?\s*\S+(?:\s*\(\d+%\)))*)', full_text)
    if m:
        data.reviewer = m.group(1)
    m = re.search(r'사후관리[:\s]*(\S+(?:\s*\(\d+%\))?)', full_text)
    if m:
        data.post_manager = m.group(1)


# ━━━━━━━━━━━━━━━ 별첨2 + 인증 ━━━━━━━━━━━━━━━

def _extract_appendix1_summary(tables, data: InvestmentReportData):
    """별첨1 주요투자조건 요약서에서 투자조건을 우선 추출."""
    for table in tables:
        first_row_text = " ".join(c.text.strip() for c in table.rows[0].cells) if table.rows else ""
        if '투자구분' not in first_row_text and '투자형태' not in first_row_text:
            continue
        # 주당 인수가격이 있는지 확인 (별첨1 특징)
        all_text = " ".join(c.text.strip() for row in table.rows for c in row.cells)
        if '주당 인수가격' not in all_text and '인수가격' not in all_text:
            continue

        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = " ".join(cells)

            # 투자구분 (신규발행/구주)
            if '투자구분' in row_text:
                for i, c in enumerate(cells):
                    if '투자구분' in c and i + 1 < len(cells):
                        data.investment_type = cells[i + 1]
                        break

            # 투자형태 (상환전환우선주 등) - 별첨1 우선
            if '투자형태' in row_text and not data.stock_type:
                for i, c in enumerate(cells):
                    if '투자형태' in c and i + 1 < len(cells):
                        val = cells[i + 1].strip()
                        # RCPS, CPS 등 약어도 풀네임으로 변환
                        if val == 'RCPS' or '상환전환' in val:
                            data.stock_type = '상환전환우선주'
                        elif val == 'CPS' or '전환우선' in val:
                            data.stock_type = '전환우선주'
                        elif val:
                            data.stock_type = val
                        break

            # 주당 인수가격 → 투자단가
            if ('주당 인수가격' in row_text or '인수가격' in row_text) and not data.issue_price:
                for i, c in enumerate(cells):
                    if '인수가격' in c and i + 1 < len(cells):
                        data.issue_price = cells[i + 1].strip()
                        break

            # 인수 주식수
            if '인수 주식수' in row_text and not data.total_shares:
                for i, c in enumerate(cells):
                    if '인수 주식수' in c and i + 1 < len(cells):
                        data.total_shares = cells[i + 1].strip()
                        break

            # 당사 투자금액
            if '당사 투자금액' in row_text and not data.investment_amount:
                for i, c in enumerate(cells):
                    if '당사 투자금액' in c and i + 1 < len(cells):
                        data.investment_amount = cells[i + 1].strip()
                        break

            # 기업가치
            if 'Pre' in row_text:
                for i, c in enumerate(cells):
                    if 'Pre' in c and i + 1 < len(cells):
                        data.pre_value = cells[i + 1].strip()
                    if 'Post' in c and i + 1 < len(cells):
                        data.post_value = cells[i + 1].strip()

            # 주요조건
            if '주요조건' in row_text:
                for c in reversed(cells):
                    if c and '주요조건' not in c and len(c) > 10:
                        # 주요조건 텍스트에서 세부 내용 파싱
                        _parse_conditions_text(c, data)
                        break

        # 동반투자 테이블 (별첨1 바로 다음에 있는 경우가 많음)
        break  # 첫 번째 매칭 테이블만 처리


def _parse_conditions_text(text: str, data: InvestmentReportData):
    """주요조건 텍스트에서 존속기간, 상환조건, 전환조건 등을 파싱."""
    # 존속기간
    m = re.search(r'존속기간\s*(\d+)\s*년', text)
    if m and not data.duration:
        data.duration = m.group(1) + "년"

    # 상환조건
    m = re.search(r'(\d+)\s*년.*?(?:후|경과).*?상환', text)
    if m and not data.redemption_terms:
        years = m.group(1)
        ytm = re.search(r'YTM\s*(\d+)\s*%', text)
        rate = ytm.group(1) if ytm else ""
        data.redemption_terms = f"{years}년후부터 상환청구 가능" + (f", 연복리 {rate}%" if rate else "")

    # Refixing
    m = re.search(r'IPO/M&?A.*?(\d+)\s*%', text)
    if m and not data.refixing_terms:
        data.refixing_terms = f"IPO/M&A {m.group(1)}%"


def _extract_co_investors_table(tables, data: InvestmentReportData):
    """동반투자 테이블에서 투자기관/금액/형태 추출."""
    for table in tables:
        first_text = " ".join(c.text.strip() for c in table.rows[0].cells) if table.rows else ""
        if '투자기관' not in first_text:
            continue

        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if '합계' in " ".join(cells) or '합 계' in " ".join(cells):
                break
            if len(cells) >= 3:
                names = cells[0].split('\n') if cells[0] else []
                amounts = cells[1].split('\n') if len(cells) > 1 else []
                # 케이런 자기 건은 제외
                for k, name in enumerate(names):
                    name = name.strip()
                    if not name or '케이런' in name:
                        continue
                    amt = amounts[k].strip() if k < len(amounts) else ""
                    if amt:
                        data.co_investors.append((name, amt, ""))
        break


def _find_yn_value(cells: list) -> str:
    """셀 목록에서 해당/미해당/O/X 값을 찾아 반환."""
    valid = {'해당', '미해당', 'O', 'X', '가능', '불가', '적합', '부적합', '아님'}
    for c in cells:
        c_clean = c.strip()
        if c_clean in valid:
            return c_clean
    return ""


def _extract_appendix2(tables, data: InvestmentReportData):
    """별첨2 투자재원검토보고서에서 주목적투자, 투자구분 등을 추출."""
    for table in tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = " ".join(cells)

            # 투자구분 (신규발행/구주)
            if '투자구분' in row_text and not data.investment_type:
                for c in cells:
                    if c and '투자구분' not in c:
                        data.investment_type = c
                        break

            # 후속투자 여부
            if '후속투자' in row_text and not data.is_follow_up:
                for c in cells:
                    if c and '후속투자' not in c and '항' not in c:
                        data.is_follow_up = c
                        break

            # 주목적 - 국토교통분야
            if '국토교통' in row_text and not data.purpose_transport:
                data.purpose_transport = _find_yn_value(cells)

            # 주목적 - 혁신성장 모빌리티
            if '모빌리티' in row_text and not data.purpose_mobility:
                data.purpose_mobility = _find_yn_value(cells)

            # 주목적 - 남부권 전략산업
            if '남부권' in row_text and not data.purpose_south:
                data.purpose_south = _find_yn_value(cells)

            # 주목적 - TCB
            if ('TCB' in row_text or 'Ti-' in row_text or 'TI-' in row_text) and '투자대상' not in row_text:
                if not data.purpose_tcb:
                    for c in cells:
                        c_clean = c.strip()
                        if c_clean in ('해당', '미해당', 'O', 'X', '가능', '불가'):
                            data.purpose_tcb = c_clean
                            break
                # TCB 상세 - "TI-3 등급(2025.8.28 발급)" 형태
                # 반드시 "TI-숫자" + "발급" or "등급" 패턴이어야 함
                for c in cells:
                    m_tcb = re.search(r'TI-(\d+)\s*등급', c)
                    if m_tcb:
                        data.purpose_tcb_detail = c.strip()
                        break

            # 표준산업분류코드
            if '표준산업' in row_text or ('주요사업' in row_text and not data.industry_code):
                for c in cells:
                    m = re.search(r'\(([A-Z]\d{4,5})\)', c)
                    if m:
                        data.industry_code = m.group(1)
                        break

    # industry_code fallback: 회사개요 테이블에서
    if not data.industry_code:
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    m = re.search(r'\(([A-Z]\d{4,5})\)', cell.text)
                    if m:
                        data.industry_code = m.group(1)
                        return


def _extract_certifications(full_text: str, data: InvestmentReportData):
    """벤처기업/이노비즈 인증 정보를 텍스트에서 추출."""
    # 벤처기업 인증
    if re.search(r'벤처기업.*?인증|벤처기업확인서|벤처.*?확인', full_text):
        data.is_venture = "Y"
        m = re.search(r'벤처.*?유효기간.*?(\d{4}[\.\-]\d{1,2}[\.\-]\d{1,2})', full_text)
        if m:
            data.venture_expiry = m.group(1)

    # 이노비즈/메인비즈
    if re.search(r'[Ii]nno-?[Bb]iz|이노비즈|기술혁신형', full_text):
        data.is_innobiz = "Y"
        m = re.search(r'[Ii]nno.*?유효기간.*?(\d{4}[\.\-]\d{1,2}[\.\-]\d{1,2})', full_text)
        if m:
            data.innobiz_expiry = m.group(1)
    if re.search(r'[Mm]ain-?[Bb]iz|메인비즈|경영혁신형', full_text):
        data.is_innobiz = data.is_innobiz or "Y"
