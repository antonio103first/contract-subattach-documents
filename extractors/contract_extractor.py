"""투자계약서에서 데이터를 추출하는 모듈.
어떤 양식이든 키워드 기반으로 자동 탐색한다. DOCX/PDF/HWP 지원."""
import re
from dataclasses import dataclass, field
from docx import Document


def _read_hwp_text(filepath: str) -> str:
    """HWP v5.0 파일에서 PrvText 스트림으로 텍스트를 추출."""
    import olefile
    ole = olefile.OleFileIO(filepath)
    try:
        prvtext = ole.openstream('PrvText').read()
        text = prvtext.decode('utf-16-le', errors='replace')
    finally:
        ole.close()
    return text


@dataclass
class InvestmentContractData:
    company_name: str = ""
    representative: str = ""
    address: str = ""
    interested_party: str = ""

    stock_type: str = ""
    total_shares: str = ""
    par_value: str = ""
    issue_price: str = ""
    total_investment: str = ""
    payment_date: str = ""
    contract_date: str = ""

    duration: str = ""
    redemption_terms: str = ""
    conversion_terms: str = ""
    refixing_terms: str = ""
    other_terms: str = ""

    fund_usage: str = ""

    article_fund_usage: str = ""
    article_consent: str = ""
    article_buyback: str = ""
    article_damages: str = ""
    article_delay_penalty: str = ""

    buyback_rate: str = ""
    penalty_rate: str = ""
    delay_rate: str = ""
    redemption_rate: str = ""

    warnings: list = field(default_factory=list)


# ── 지역 리스트 (주소 추출용) ──
_REGIONS = (
    '서울|경기|인천|부산|대구|광주|대전|울산|세종|강원|충북|충남'
    '|전북|전남|경북|경남|제주|서울특별시|부산광역시|대구광역시'
    '|인천광역시|광주광역시|대전광역시|울산광역시|대전시|부산시'
)


def extract_contract_data(filepath: str) -> InvestmentContractData:
    ext = filepath.lower().rsplit('.', 1)[-1] if '.' in filepath else ''
    if ext == 'pdf':
        from extractors.pdf_extractor import extract_text_from_pdf
        full_text = extract_text_from_pdf(filepath)
        paragraphs = [p.strip() for p in full_text.split('\n') if p.strip()]
        doc = None
    elif ext == 'hwp':
        full_text = _read_hwp_text(filepath)
        paragraphs = [p.strip() for p in full_text.split('\n') if p.strip()]
        doc = None
    else:
        doc = Document(filepath)
        paragraphs = [p.text.strip() for p in doc.paragraphs]
        full_text = "\n".join(paragraphs)

    data = InvestmentContractData()
    _extract_parties(full_text, data)
    _extract_share_info(full_text, data)
    _extract_contract_date(full_text, data)
    _extract_article_numbers(paragraphs, full_text, data)
    _extract_penalty_rates(full_text, data)
    _extract_preferred_stock_terms(full_text, data)
    _extract_fund_usage(full_text, doc, data)
    return data


# ━━━━━━━━━━━━━━━━━━ 당사자 정보 ━━━━━━━━━━━━━━━━━━

def _extract_parties(full_text: str, data: InvestmentContractData):
    """회사명, 대표이사, 주소, 이해관계인을 추출."""

    # ── 섹션 분리: "1. 투자자 / 2. 피투자자(회사) / 3. 이해관계인" ──
    # 다양한 표현: 투자기업, 피투자자, 회사, 발행회사
    company_section = ""
    for pat in [
        r'(2\.\s*(?:투자기업|피투자자|회사).*?)(?=3\.\s*이해관계인)',
        r'(발행회사\s*[:：].*?)(?=투\s*자\s*자|이해관계인)',
    ]:
        m = re.search(pat, full_text, re.DOTALL)
        if m:
            company_section = m.group(1)
            break

    # ── 회사명 ── (본문 앞부분 우선, 별지 안의 섹션은 후순위)
    search_areas = [full_text[:5000]]
    if company_section:
        search_areas.append(company_section)

    for area in search_areas:
        if data.company_name:
            break
        for pat in [
            r'(?:회사명|발행회사|발행인|피투자자|투자기업)\s*[:：]\s*\n?\s*((?:주식회사|㈜|\(주\))\s*\S+)',
            r'(?:^회사)\s*[:：]?\s*\n\s*((?:주식회사|㈜|\(주\))\s*\S+)',  # "회사\n주식회사 XXX"
            r'(?:회사)\s*[:：]\s*\n?\s*((?:주식회사|㈜|\(주\))\s*\S+)',
            r'"회사"의\s*상호\s*[:：]\s*((?:주식회사|㈜|\(주\))\s*\S+)',
        ]:
            m = re.search(pat, area, re.MULTILINE)
            if m:
                name = m.group(1).strip().rstrip('"\')')
                if '케이런' not in name and '조합' not in name:
                    data.company_name = name
                    break

    # ── 대표이사 (회사의 대표, 투자자 대표 아님) ──
    search_area = company_section or full_text[:3000]
    # 이해관계인 섹션에서도 시도
    interested_m = re.search(r'(3\.\s*이해관계인.*?)(?=제\d|다\s*음|\n\n\n)', full_text, re.DOTALL)
    interested_section = interested_m.group(1) if interested_m else ""

    for area in [search_area, interested_section]:
        if data.representative:
            break
        for pat in [
            r'대표이사\s*[:：]?\s*([가-힣]{2,4})',
            r'대표이사\s+([가-힣]\s*[가-힣]\s*[가-힣])',  # 공백 포함
        ]:
            m = re.search(pat, area)
            if m:
                data.representative = m.group(1).replace(' ', '').strip()
                break

    # ── 주소 (회사 주소) ──
    # "회사" 라벨 다음의 주소를 추출 (투자자 주소 제외)
    # 패턴: "회사 : 주식회사 XXX\n주소 : ..."
    m = re.search(
        r'(?:"?회사"?|피투자자)\s*[:：].*?(?:주식회사|㈜).*?\n\s*주소\s*[:：]\s*(' + _REGIONS + r')([^\n]{5,80})',
        full_text[:5000], re.DOTALL
    )
    if m:
        data.address = (m.group(1) + m.group(2)).strip()

    if not data.address:
        for area in [company_section, full_text[:5000]]:
            if data.address:
                break
            m = re.search(r'주소\s*[:：]\s*(' + _REGIONS + r')([^\n]{5,80})', area)
            if m and '테헤란' not in m.group(0) and '송강빌딩' not in m.group(0) and '강남구' not in m.group(0):
                data.address = (m.group(1) + m.group(2)).strip()
                break

    # ── 이해관계인 ──
    m = re.search(r'이해관계인.*?대표이사\s*([가-힣]{2,4})', full_text)
    if m:
        data.interested_party = m.group(1).replace(' ', '')
    elif data.representative:
        data.interested_party = data.representative


# ━━━━━━━━━━━━━━━━━━ 신주 발행 정보 ━━━━━━━━━━━━━━━━━━

def _extract_share_info(full_text: str, data: InvestmentContractData):
    """신주 종류, 수량, 가격 등을 추출. 조문 번호에 무관."""

    # ── 신주 종류 ──
    for pat in [
        r'종류\s*(?:와\s*수|및\s*내용)\s*[:：]?\s*(?:기명식\s*)?(\S+우선주)\S*\s',
        r'본건 신주의 종류\s*\n\s*(.+)',
        r'(상환전환우선주식?|전환우선주식?|상환우선주식?)',
    ]:
        m = re.search(pat, full_text)
        if m and not data.stock_type:
            raw = m.group(1).strip()
            raw = re.sub(r'\(이하.*', '', raw).strip()
            raw = re.sub(r'주식$', '주', raw)
            data.stock_type = raw
            break
    if not data.stock_type:
        m = re.search(r'(상환전환우선주|전환우선주|상환우선주|보통주)', full_text[:2000])
        if m:
            data.stock_type = m.group(1)

    # ── 발행 주식수 ──
    for pat in [
        r'(?:본건 신주의 발행 총수|종류와 수|"본 주식"의 종류와 수)\s*[:：]?\s*(?:기명식\s*\S+\s*)?([\d,]+)\s*주',
        r'(?:발행 총수|발행총수)\s*[:：]\s*([\d,]+)\s*주',
    ]:
        m = re.search(pat, full_text)
        if m and not data.total_shares:
            data.total_shares = m.group(1)
            break

    # ── 액면가 ──
    for pat in [
        r'(?:1주당 액면가액|1주의 금액|액면가)\s*(?:\([^)]*\))?\s*[:：]\s*금?\s*([\d,]+)\s*원',
    ]:
        m = re.search(pat, full_text)
        if m and not data.par_value:
            data.par_value = m.group(1)
            break

    # ── 1주당 발행가액 ──
    for pat in [
        r'1주당 발행가액\s*(?:\([^)]*\))?\s*[:：]\s*금?\s*([\d,]+)\s*원',
        r'발행가액\s*(?:\([^)]*\))?\s*[:：]\s*금?\s*(?:\s*)([\d,]+)\s*원',
    ]:
        m = re.search(pat, full_text)
        if m and not data.issue_price:
            data.issue_price = m.group(1)
            break

    # ── 총 인수가액/인수대금 ──
    for pat in [
        r'총 인수(?:가액|대금)\s*[:：].*?([\d,]{7,})\s*원',
        r'\\([\d,]{7,})\s*원\)',
        r'인수(?:가액|대금)\s*[:：].*?([\d,]{7,})\s*원',
    ]:
        m = re.search(pat, full_text)
        if m and not data.total_investment:
            data.total_investment = m.group(1)
            break

    # ── 납입기일 ──
    m = re.search(r'납입기일\s*[:：]\s*(\d{4})년\s*(\d{1,2})월\s*\[?\s*(\d{1,2})?\s*\]?\s*일', full_text)
    if m:
        day = m.group(3) or ""
        data.payment_date = f"{m.group(1)}년 {m.group(2)}월 {day}일".strip()


def _extract_contract_date(full_text: str, data: InvestmentContractData):
    for pat in [
        r'(\d{4})년\s*(\d{1,2})월\s*(\d{1,2})일.*?본 계약 체결일',
        r'본.*?계약.*?(\d{4})년\s*(\d{1,2})월\s*(\d{1,2})일',
        r'(\d{4})\.\s*(\d{1,2})\.\s*(\d{1,2})',
    ]:
        m = re.search(pat, full_text[:2000])
        if m:
            data.contract_date = f"{m.group(1)}년 {m.group(2)}월 {m.group(3)}일"
            break


# ━━━━━━━━━━━━━━━━━━ 조문 번호 탐색 ━━━━━━━━━━━━━━━━━━

def _extract_article_numbers(paragraphs: list, full_text: str, data: InvestmentContractData):
    """키워드 기반으로 조문 번호를 자동 탐색."""
    keyword_map = {
        'article_fund_usage': ['투자금의 용도', '투자금 사용'],
        'article_consent': ['동의권', '협의권', '경영사항'],
        'article_buyback': ['주식매수청구권'],
        'article_damages': ['위약벌', '손해배상'],
        'article_delay_penalty': ['지연배상금', '지연손해금'],
    }

    # 방법 1: 줄 시작에서 "제N조 제목" 패턴
    article_pattern = re.compile(r'^제\s*(\d+)\s*조\s*[\(\s]*(.+?)[\)\s]*$', re.MULTILINE)
    for m in article_pattern.finditer(full_text):
        num = m.group(1)
        title = m.group(2).strip()
        # 너무 긴 제목은 본문 참조이므로 제외
        if len(title) > 50:
            continue
        for field_name, keywords in keyword_map.items():
            if getattr(data, field_name):
                continue
            for kw in keywords:
                if kw in title:
                    setattr(data, field_name, num)
                    break

    # 방법 2: 소제목(단독 줄) + 본문 구조
    # 예: "주식매수청구권\n\n다음 각 호의..." → 장/절 번호로 추정
    heading_keywords = {
        'article_fund_usage': ['투자금의 용도 및 제한', '투자금의 용도'],
        'article_consent': ['경영사항에 대한 동의권', '동의권 및 협의권'],
        'article_buyback': ['주식매수청구권'],
        'article_damages': ['위약벌 및 손해배상', '위약벌', '손해배상'],
        'article_delay_penalty': ['지연배상금', '지연손해금'],
    }

    for field_name, keywords in heading_keywords.items():
        if getattr(data, field_name):
            continue
        for kw in keywords:
            for pat in [
                # 키워드가 소제목이고 이후 본문에 조문 참조
                re.compile(kw + r'[^\n]*\n[^\n]*?제\s*(\d+)\s*조'),
                # 본문 내에서 "제N조(키워드)"
                re.compile(r'제\s*(\d+)\s*조\s*[\(（]' + re.escape(kw)),
                # 본문에서 "제N조 키워드"
                re.compile(r'제\s*(\d+)\s*조\s+' + re.escape(kw)),
            ]:
                m = pat.search(full_text)
                if m:
                    num = m.group(1)
                    # 상법 조문(300번대 이상) 제외
                    if int(num) < 200:
                        setattr(data, field_name, num)
                        break
            if getattr(data, field_name):
                break

    # 방법 3: 조문 번호가 없는 구조 → 장(chapter) 번호 + 키워드 본문 위치로 추정
    # "제5장 계약 위반에 대한 책임\n\n주식매수청구권\n\n..." 같은 구조
    # 이 경우 본문 참조에서 자기 조문을 언급하는 패턴 탐색
    for field_name, keywords in heading_keywords.items():
        if getattr(data, field_name):
            continue
        for kw in keywords:
            idx = full_text.find(kw)
            if idx < 0:
                continue
            # 키워드 이후 2000자 범위에서 "본조" 또는 자기참조 패턴 탐색
            after = full_text[idx:idx+3000]
            # "제N조에 따른" 또는 "제N조의" 패턴에서 N 추출
            refs = re.findall(r'제\s*(\d+)\s*조', after[:500])
            # 가장 자주 나오는 조문 번호가 해당 조문일 가능성 높음
            if refs:
                from collections import Counter
                counts = Counter(refs)
                # 상법 조문 제외
                valid = [(n, c) for n, c in counts.items() if int(n) < 100]
                if valid:
                    most_common = max(valid, key=lambda x: x[1])
                    setattr(data, field_name, most_common[0])
                    break


# ━━━━━━━━━━━━━━━━━━ 의무불이행 이자율 ━━━━━━━━━━━━━━━━━━

def _extract_penalty_rates(full_text: str, data: InvestmentContractData):
    # 주식매수청구권 이자율
    section = _find_section(full_text, '주식매수청구권')
    if section:
        m = re.search(r'연\s*(\d+)\s*%', section)
        if m:
            data.buyback_rate = m.group(1)

    # 위약벌
    for kw in ['위약벌', '손해배상']:
        section = _find_section(full_text, kw)
        if section and not data.penalty_rate:
            m = re.search(r'투자금의?\s*(\d+)\s*%', section)
            if m:
                data.penalty_rate = m.group(1)
            else:
                m = re.search(r'(\d+)\s*%.*?위약벌', section)
                if m:
                    data.penalty_rate = m.group(1)

    # 지연배상금
    for kw in ['지연배상금', '지연손해금']:
        section = _find_section(full_text, kw)
        if section and not data.delay_rate:
            m = re.search(r'연\s*(\d+)\s*%', section)
            if m:
                data.delay_rate = m.group(1)

    # 상환 이자율
    for pat in [
        r'상환.*?연\s*복리\s*(\d+)\s*%',
        r'연\s*복리\s*(\d+)\s*%.*?상환',
        r'YTM\s*(\d+)\s*%',
    ]:
        m = re.search(pat, full_text)
        if m and not data.redemption_rate:
            data.redemption_rate = m.group(1)
            break


# ━━━━━━━━━━━━━━━━━━ 우선주 조건 ━━━━━━━━━━━━━━━━━━

def _extract_preferred_stock_terms(full_text: str, data: InvestmentContractData):
    # 존속기간
    m = re.search(r'존속기간.*?(\d+)\s*년', full_text)
    if m:
        data.duration = f"{m.group(1)}년"

    # 상환조건
    for pat in [
        r'효력발생일로부터\s*(\d+)\s*년이?\s*경과.*?상환',
        r'발행일로?부터\s*(\d+)\s*년이?\s*(?:경과|후).*?상환',
        r'(\d+)\s*년.*?경과.*?상환청구',
    ]:
        m = re.search(pat, full_text)
        if m and not data.redemption_terms:
            years = m.group(1)
            rate = data.redemption_rate
            data.redemption_terms = (
                f"투자 {years}년후부터 연복리 {rate}%로 상환청구가능" if rate
                else f"투자 {years}년후부터 상환청구가능"
            )
            break

    # 전환조건
    parts = []
    if re.search(r'효력발생일.*?(?:존속기간|만료).*?전환', full_text):
        parts.append("발행일 익일부터 만기까지")
    if re.search(r'존속기간.*?만료.*?보통주.*?전환', full_text):
        parts.append("존속기간 이후 보통주 자동 전환")
    m = re.search(r'(?:종류주식|우선주)\s*1주.*?보통주.*?(\d+)\s*주', full_text)
    if m:
        parts.append(f"우선주 1주당 보통주 {m.group(1)}주로 전환")
    data.conversion_terms = ", ".join(parts) if parts else ""

    # Refixing
    rparts = []
    m = re.search(r'공모단가.*?(\d+)\s*%', full_text)
    if m:
        rparts.append(f"IPO/M&A {m.group(1)}%")
    if re.search(r'전환가[격액]보다 낮은', full_text):
        rparts.append("투자단가 이하 유상증자 등")
    data.refixing_terms = "Refixing: " + ", ".join(rparts) if rparts else ""

    # 기타 권리
    oparts = []
    if re.search(r'우선매수권', full_text):
        oparts.append("우선매수권")
    if re.search(r'공동매도참여권|공동매도권|Tag.?Along', full_text):
        oparts.append("공동매도참여권")
    if re.search(r'Drag.?Along|동반매도청구권', full_text):
        oparts.append("동반매도청구권")
    data.other_terms = ", ".join(oparts) if oparts else ""


# ━━━━━━━━━━━━━━━━━━ 투자금 사용용도 ━━━━━━━━━━━━━━━━━━

def _extract_fund_usage(full_text: str, doc, data: InvestmentContractData):
    # 테이블에서 찾기
    if doc:
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                joined = " ".join(cells)
                if ('사용용도' in joined or '사용목적' in joined) and '위반' not in joined:
                    for ct in cells:
                        if ct and '사용용도' not in ct and '사용목적' not in ct and len(ct) > 5:
                            data.fund_usage = ct
                            return

    # 텍스트에서 찾기
    m = re.search(r'별지\s*\d+.*?투자금.*?사용.*?용도\s*\n(.+)', full_text)
    if m:
        data.fund_usage = m.group(1).strip()


# ━━━━━━━━━━━━━━━━━━ 유틸 ━━━━━━━━━━━━━━━━━━

def _find_section(full_text: str, keyword: str) -> str:
    """키워드를 포함하는 조문의 본문 텍스트를 반환."""
    # "제N조(...keyword...)" 또는 "제N조 ...keyword..."
    for pat in [
        re.compile(r'제\s*\d+\s*조\s*[\(（]?' + r'[^)\n]*' + re.escape(keyword) + r'.*?\n', re.DOTALL),
        re.compile(r'제\s*\d+\s*조\s+[^\n]*' + re.escape(keyword) + r'.*?\n', re.DOTALL),
    ]:
        m = pat.search(full_text)
        if m:
            start = m.start()
            rest = full_text[m.end():]
            next_art = re.search(r'\n제\s*\d+\s*조[\s(（]', rest)
            end = m.end() + next_art.start() if next_art else start + 5000
            return full_text[start:end]
    return ""
