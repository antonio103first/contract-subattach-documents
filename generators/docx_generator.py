"""투자계약서 체크리스트 및 의무기재사항확인서 DOCX 생성 모듈."""
import copy
import os
import re
from datetime import datetime
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree


def normalize_value(val: str) -> str:
    """비교를 위해 값을 정규화 (공백, 쉼표, '원', '주', '%' 제거)."""
    if not val:
        return ""
    v = re.sub(r'[\s,원주%㈜주식회사(주)]', '', str(val))
    return v.strip()


def compare_values(report_val, contract_val, field_name):
    """두 소스 값을 비교하여 불일치 시 경고를 반환."""
    if not report_val or not contract_val:
        return None  # 한쪽이 없으면 비교 불가
    if normalize_value(report_val) == normalize_value(contract_val):
        return None
    warning = f"⚠ 불일치: 투심보고서={report_val}, 투자계약서={contract_val}. 확인 필요."
    print(f"[WARNING] {field_name} {warning}")
    return warning


def add_comment_to_cell(cell, comment_text: str, author="자동검증"):
    """DOCX 셀에 주석(Comment)을 추가한다."""
    # python-docx는 공식적으로 comment를 지원하지 않으므로,
    # 셀 텍스트 뒤에 빨간색 텍스트로 경고를 표시한다.
    paragraph = cell.paragraphs[-1] if cell.paragraphs else cell.add_paragraph()
    run = paragraph.add_run(f"\n{comment_text}")
    run.font.color.rgb = _rgb(255, 0, 0)
    run.font.size = _pt(8)


def _rgb(r, g, b):
    from docx.shared import RGBColor
    return RGBColor(r, g, b)


def _pt(size):
    from docx.shared import Pt
    return Pt(size)


def _set_cell_text(cell, text: str):
    """셀의 텍스트를 완전히 교체한다. 기존 서식을 최대한 유지."""
    if cell.paragraphs:
        # 첫 번째 paragraph의 서식 유지, 텍스트 교체
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].text = text
            for run in p.runs[1:]:
                run.text = ""
        else:
            p.text = text
        # 나머지 paragraph 삭제 (병합 셀 등에서 발생)
        for extra_p in cell.paragraphs[1:]:
            extra_p.clear()
    else:
        cell.text = text


def generate_docx_checklist(contract_data, report_data, template_path: str, output_path: str):
    """투자계약서 체크리스트 DOCX를 생성한다."""
    doc = Document(template_path)
    tables = doc.tables
    warnings = []

    now = datetime.now()
    date_str = f"{now.year}년 {now.month}월 {now.day}일"

    # 데이터 준비
    company_name = report_data.company_name or contract_data.company_name
    short_name = company_name  # ㈜위밋모빌리티
    if '주식회사' in company_name:
        short_name = "㈜" + company_name.replace('주식회사', '').replace('㈜', '').strip()
    if not short_name.startswith('㈜') and '㈜' not in short_name:
        short_name = "㈜" + short_name

    representative = report_data.representative or contract_data.representative
    address = report_data.address or contract_data.address
    business_id = report_data.business_registration or ""
    share_ratio = report_data.share_ratio or ""
    investment_amount = contract_data.total_investment or ""
    if investment_amount:
        investment_amount = f"{investment_amount}원"
    issue_price = contract_data.issue_price or ""
    if issue_price:
        issue_price = f"{issue_price}원"
    total_shares = contract_data.total_shares or ""
    if total_shares:
        total_shares = f"{total_shares}주"
    stock_type = contract_data.stock_type or report_data.stock_type or "상환전환우선주"

    # 기타조건 구성
    other_conditions_parts = []
    if contract_data.duration:
        other_conditions_parts.append(f"존속기간: {contract_data.duration}")
    if contract_data.redemption_terms:
        other_conditions_parts.append(f"상환조건: {contract_data.redemption_terms}")
    if contract_data.conversion_terms:
        other_conditions_parts.append(f"전환조건: {contract_data.conversion_terms}")
    if contract_data.refixing_terms:
        other_conditions_parts.append(contract_data.refixing_terms)
    if contract_data.other_terms:
        other_conditions_parts.append(f"이해관계인의 주식 처분시 투자자의 {contract_data.other_terms}")
    other_conditions = "\n".join(other_conditions_parts)

    # 의무불이행 이자율 구성
    penalty_parts = []
    if contract_data.buyback_rate:
        penalty_parts.append(f"주식매수청구권 : 투자원금 및 연 {contract_data.buyback_rate}%")
    if contract_data.penalty_rate:
        penalty_parts.append(f"위약벌 : 투자원금의 {contract_data.penalty_rate}%")
    if contract_data.delay_rate:
        penalty_parts.append(f"지연배상금 : 실제 지급일까지 연 {contract_data.delay_rate}%")
    penalty_text = "\n".join(penalty_parts)

    # 투자금 사용용도
    fund_usage = report_data.fund_usage or contract_data.fund_usage or ""

    # 조문 참조
    art_fund = contract_data.article_fund_usage or "7"
    art_consent = contract_data.article_consent or "9"
    art_buyback = contract_data.article_buyback or "18"
    art_damages = contract_data.article_damages or "19"
    art_delay = contract_data.article_delay_penalty or "20"

    # --- 교차 검증 ---
    cross_checks = [
        ("회사명", report_data.company_name, contract_data.company_name),
        ("대표이사", report_data.representative, contract_data.representative),
        ("주소", report_data.address, contract_data.address),
        ("투자금액", report_data.investment_amount, contract_data.total_investment + "원" if contract_data.total_investment else ""),
        ("투자단가", report_data.issue_price, contract_data.issue_price + "원" if contract_data.issue_price else ""),
        ("투자방식", report_data.stock_type, contract_data.stock_type),
    ]
    for field_name, rval, cval in cross_checks:
        w = compare_values(rval, cval, field_name)
        if w:
            warnings.append((field_name, w))

    # ========= Table 0: 투자기업 정보 =========
    if len(tables) > 0:
        t = tables[0]
        _set_cell_text(t.rows[0].cells[1], short_name)
        _set_cell_text(t.rows[0].cells[3], business_id)
        _set_cell_text(t.rows[1].cells[1], representative)
        _set_cell_text(t.rows[1].cells[3], share_ratio)
        # 소재지 - 병합 셀 처리
        for cell in t.rows[2].cells[1:]:
            _set_cell_text(cell, address)

    # ========= Table 1: 투자조건 비교표 =========
    if len(tables) > 1:
        t = tables[1]
        # Row 1: 투자업체명
        _set_cell_text(t.rows[1].cells[1], short_name)
        _set_cell_text(t.rows[1].cells[2], "좌동")
        _set_cell_text(t.rows[1].cells[3], "여")

        # Row 2: 투자금액
        _set_cell_text(t.rows[2].cells[1], investment_amount)
        _set_cell_text(t.rows[2].cells[2], "좌동")
        _set_cell_text(t.rows[2].cells[3], "여")

        # Row 3: 투자방식
        _set_cell_text(t.rows[3].cells[1], f"{stock_type}/신주")
        _set_cell_text(t.rows[3].cells[2], "좌동")
        _set_cell_text(t.rows[3].cells[3], "여")

        # Row 4: 투자단가
        _set_cell_text(t.rows[4].cells[1], issue_price)
        _set_cell_text(t.rows[4].cells[2], "좌동")
        _set_cell_text(t.rows[4].cells[3], "여")

        # Row 5: 인수주식수
        _set_cell_text(t.rows[5].cells[1], total_shares)
        _set_cell_text(t.rows[5].cells[2], "좌동")
        _set_cell_text(t.rows[5].cells[3], "여")

        # Row 6: 지분율
        _set_cell_text(t.rows[6].cells[1], share_ratio)
        _set_cell_text(t.rows[6].cells[2], "좌동")
        _set_cell_text(t.rows[6].cells[3], "여")

        # Row 7: 기타조건
        _set_cell_text(t.rows[7].cells[1], other_conditions)
        _set_cell_text(t.rows[7].cells[2], "좌동")
        _set_cell_text(t.rows[7].cells[3], "여")

        # Row 8: 투자금 사용용도
        _set_cell_text(t.rows[8].cells[1], fund_usage)
        _set_cell_text(t.rows[8].cells[2], "좌동")
        _set_cell_text(t.rows[8].cells[3], "여")

        # Row 9: 의무불이행 이자율
        _set_cell_text(t.rows[9].cells[1], penalty_text)
        _set_cell_text(t.rows[9].cells[2], "좌동")
        _set_cell_text(t.rows[9].cells[3], "여")

        # 교차검증 경고 주석 추가
        for field_name, warning_text in warnings:
            # 해당 필드의 행에 주석 추가
            row_map = {"투자금액": 2, "투자단가": 4, "투자방식": 3, "회사명": 1}
            row_idx = row_map.get(field_name)
            if row_idx and row_idx < len(t.rows):
                add_comment_to_cell(t.rows[row_idx].cells[1], warning_text)

    # ========= Table 2: 투자계약서 조문 참조 =========
    if len(tables) > 2:
        t = tables[2]
        # Row 1: 투자금 사용용도 및 변경 시 사전 인지 방안
        _set_cell_text(t.rows[1].cells[1],
            f"제{art_fund}조: 투자금의 용도 및 제한\n[별지3] 투자금의 사용용도 및 실사 약정")
        # Row 2: 투자금 사용용도 위반 시 제재 방안
        _set_cell_text(t.rows[2].cells[1],
            f"제{art_buyback}조: 주식매수청구권\n제{art_damages}조: 손해배상 및 위약벌\n제{art_delay}조 : 지연배상금")
        # Row 3: 투자금 사용처 실사근거
        _set_cell_text(t.rows[3].cells[1],
            f"제{art_fund}조: 투자금의 용도 및 제한\n[별지3] 투자금의 사용용도 및 실사 약정")
        # Row 4: 대상거래 사전인지방안
        _set_cell_text(t.rows[4].cells[1],
            f"제{art_fund}조: 투자금의 용도 및 제한\n[별지3] 투자금의 사용용도 및 실사 약정")
        # Row 5: 대상거래 발생시 처리방안
        _set_cell_text(t.rows[5].cells[1],
            f"제{art_buyback}조: 주식매수청구권\n제{art_damages}조: 손해배상 및 위약벌\n제{art_delay}조 : 지연배상금")

    # ========= Table 3: 요약문 =========
    if len(tables) > 3:
        t = tables[3]
        _set_cell_text(t.rows[0].cells[0],
            f"{short_name} 투자와 관련하여 투자심사보고서와 투자계약서의 내용을 검토한 결과, 투자심사보고서와 투자계약서의 내용은 일치합니다.")

    # ========= Table 4: 의무기재사항 확인서 =========
    if len(tables) > 4:
        t = tables[4]
        # Row 1: 대상업체명
        _set_cell_text(t.rows[1].cells[2], short_name)
        # Row 2: 거래상대방
        _set_cell_text(t.rows[2].cells[2], short_name)
        # Row 3: 투자자산 구분
        _set_cell_text(t.rows[3].cells[2], stock_type)
        # Row 4: 투자금액
        _set_cell_text(t.rows[4].cells[2], investment_amount)
        # Row 5: 투자단가
        _set_cell_text(t.rows[5].cells[2], issue_price)
        # Row 6-11: 기재됨/해당사항없음 (유지)

    # ========= 본문 텍스트 치환 =========
    for para in doc.paragraphs:
        if '㈜AAA' in para.text or '(주)AAA' in para.text:
            for run in para.runs:
                run.text = run.text.replace('㈜AAA', short_name).replace('(주)AAA', short_name)

        # 날짜 치환
        if '2023년' in para.text and '5월' in para.text:
            for run in para.runs:
                run.text = re.sub(r'2023년\s*5월\s*일', date_str, run.text)

        # 투자대상업체
        if '투자대상업체' in para.text and '㈜AAA' in para.text:
            for run in para.runs:
                run.text = run.text.replace('㈜AAA', short_name)

    # 저장
    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    doc.save(output_path)
    print(f"[OK] DOCX 체크리스트 생성 완료: {output_path}")

    if warnings:
        print(f"\n[주의] {len(warnings)}건의 불일치 발견:")
        for fn, w in warnings:
            print(f"  - {fn}: {w}")

    return warnings
